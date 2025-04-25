from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO

def create_docx(content: str) -> bytes:
    doc = Document()
    
    # Set margin
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # Add content
    for line in content.split('\n'):
        if line.strip():  # Skip empty lines
            if any(header in line.lower() for header in ['summary', 'experience', 'education', 'skills']):
                heading = doc.add_heading(level=1)
                heading.add_run(line.strip()).bold = True
                heading.paragraph_format.space_after = Pt(12)
            else:
                para = doc.add_paragraph(line.strip())
                para.paragraph_format.space_after = Pt(6)
    
    # Save to bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()